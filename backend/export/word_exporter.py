"""
backend/export/word_exporter.py
python-docx を用いた Word (.docx) レポート出力エンジン。

依存: python-docx>=1.1
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from .base import BaseExporter


class WordExporter(BaseExporter):
    """解析結果を Word (.docx) レポートに変換するエクスポータ。"""

    _ACCENT = RGBColor(0x1A, 0x3A, 0x5C)
    _HEADER_BG = "1a3a5c"

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        """表セルの背景色を設定する（python-docx は直接APIがないため XML 直接操作）。"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = self._ACCENT
        h.runs[0].font.size = Pt(16 if level == 1 else 13)

    def _add_metrics_table(self, doc: Document, metrics: dict[str, Any]) -> None:
        """評価指標テーブルを追加する。"""
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # ヘッダー行
        hdr = table.rows[0].cells
        hdr[0].text = "指標"
        hdr[1].text = "値"
        for cell in hdr:
            self._set_cell_bg(cell, self._HEADER_BG)
            for para in cell.paragraphs:
                run = para.runs[0] if para.runs else para.add_run(cell.text)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(11)

        # データ行
        for key, val in metrics.items():
            row = table.add_row().cells
            row[0].text = key
            row[1].text = f"{val:.4f}" if isinstance(val, float) else str(val)
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

    def _add_importance_table(self, doc: Document, importances: dict[str, float], top_n: int = 10) -> None:
        """特徴量重要度テーブル（上位 top_n 件）を追加する。"""
        sorted_imp = sorted(importances.items(), key=lambda x: x[1], reverse=True)[:top_n]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        hdr[0].text = "特徴量名"
        hdr[1].text = "重要度スコア"
        for cell in hdr:
            self._set_cell_bg(cell, self._HEADER_BG)
            for para in cell.paragraphs:
                run = para.runs[0] if para.runs else para.add_run(cell.text)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

        for feat, score in sorted_imp:
            row = table.add_row().cells
            row[0].text = feat
            row[1].text = f"{score:.4f}"

        doc.add_paragraph()

    def export(self, result: dict[str, Any], filename: str) -> Path:
        """解析結果を Word ファイルに変換して output_dir へ書き出す。

        Parameters
        ----------
        result : dict
            必須キー: "best_model_name", "metrics"
            任意キー: "feature_importances", "chart_paths", "ai_commentary"
        filename : str
            拡張子なしのファイル名（例: "analysis_report"）。

        Returns
        -------
        Path
            書き出した .docx ファイルの絶対パス。
        """
        out_path = self.output_dir / f"{filename}.docx"
        doc = Document()

        # ── タイトル ──
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("ChemAI ML Studio: 解析レポート")
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = self._ACCENT
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # ── 最良モデル ──
        self._add_heading(doc, "🏆 最良モデル", level=2)
        model_name = result.get("best_model_name", "N/A")
        para = doc.add_paragraph()
        para.add_run("採択モデル: ").font.bold = True
        para.add_run(model_name)

        doc.add_paragraph()

        # ── 評価指標 ──
        metrics = result.get("metrics", {})
        if metrics:
            self._add_heading(doc, "📊 評価指標", level=2)
            self._add_metrics_table(doc, metrics)

        # ── AIコメント ──
        ai_comment = result.get("ai_commentary", "")
        if ai_comment:
            self._add_heading(doc, "🤖 AIアシスタントによる考察", level=2)
            doc.add_paragraph(ai_comment)
            doc.add_paragraph()

        # ── 特徴量重要度 ──
        importances = result.get("feature_importances", {})
        if importances:
            self._add_heading(doc, "🔑 特徴量重要度 (上位10件)", level=2)
            self._add_importance_table(doc, importances)

        # ── チャート画像 ──
        chart_paths: list[Path] = [
            Path(p) for p in result.get("chart_paths", [])
            if Path(p).exists()
        ]
        if chart_paths:
            self._add_heading(doc, "📈 解析チャート", level=2)
            for cp in chart_paths:
                try:
                    doc.add_picture(str(cp), width=Inches(5.5))
                    caption = doc.add_paragraph(cp.stem)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].font.size = Pt(9)
                    caption.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    doc.add_paragraph()
                except Exception:
                    doc.add_paragraph(f"⚠️ 画像読み込み失敗: {cp.name}")

        doc.save(str(out_path))
        return out_path.resolve()
